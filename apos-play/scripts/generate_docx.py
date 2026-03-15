#!/usr/bin/env python3
"""
Convert Especificacion_de_Requerimientos_v2.md to .docx
Matches the style of the original SRS document.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
INPUT_FILE = os.path.join(PROJECT_DIR, "Especificacion_de_Requerimientos_v2.md")
OUTPUT_FILE = os.path.join(PROJECT_DIR, "Especificacion_de_Requerimientos_v2.docx")

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)
HEADER_GRAY = "D9D9D9"


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """Apply borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_run_font(run, bold=False, italic=False, size=None):
    """Configure a run's font."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    # Ensure East Asian font is also set
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)


def add_formatted_text(paragraph, text):
    """Add text to paragraph, handling **bold** and *italic* markdown."""
    # Split by bold markers first
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif '*' in part:
            # Handle italic within non-bold parts
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and not sp.startswith('**'):
                    run = paragraph.add_run(sp[1:-1])
                    set_run_font(run, italic=True)
                else:
                    if sp:
                        run = paragraph.add_run(sp)
                        set_run_font(run)
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run)


def configure_styles(doc):
    """Configure document styles to match the original."""
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = FONT_NAME
            s.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                s.font.size = Pt(18)
                s.font.bold = True
            elif level == 2:
                s.font.size = Pt(14)
                s.font.bold = True
            elif level == 3:
                s.font.size = Pt(12)
                s.font.bold = True


def add_cover_page(doc):
    """Add cover page matching original style."""
    # Spacing before title
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    # "Proyecto AposPlay"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto\nAposPlay")
    set_run_font(run, bold=True, size=Pt(26))

    doc.add_paragraph()

    # "Documento de Requisitos del Sistema"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento de Requisitos\ndel Sistema")
    set_run_font(run, bold=True, size=Pt(22))

    doc.add_paragraph()

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión 2.0")
    set_run_font(run, bold=True, italic=True, size=Pt(14))

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fecha: 14/03/2026")
    set_run_font(run, bold=True, italic=True, size=Pt(11))

    # Spacing
    for _ in range(4):
        doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    run = p.add_run("Realizado por: Ivasiuta, Mariano - Martinez, Alejandro")
    set_run_font(run)
    p = doc.add_paragraph()
    run = p.add_run("Realizado para: Cliente")
    set_run_font(run)

    doc.add_page_break()


def parse_table_rows(lines, start_idx):
    """Parse markdown table rows starting from start_idx. Returns (rows, next_idx)."""
    rows = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            idx += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        idx += 1
    return rows, idx


def add_table_to_doc(doc, rows, is_uc_table=False):
    """Add a table to the document with proper formatting."""
    if not rows:
        return

    num_cols = len(rows[0])
    # Ensure all rows have same number of columns
    for i, row in enumerate(rows):
        while len(row) < num_cols:
            row.append('')
        if len(row) > num_cols:
            rows[i] = row[:num_cols]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    # Auto-fit
    table.autofit = True

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            # Clear default paragraph
            cell.paragraphs[0].clear()

            # First row is header
            if i == 0:
                set_cell_shading(cell, HEADER_GRAY)
                add_formatted_text(cell.paragraphs[0], cell_text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            else:
                add_formatted_text(cell.paragraphs[0], cell_text)

            # Set font size for all runs
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    doc.add_paragraph()  # Space after table


def add_uc_table(doc, rows):
    """Add a UC-style 2-column table (field | value) with field column gray."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    table.autofit = True

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                continue
            cell = table.cell(i, j)
            cell.paragraphs[0].clear()
            add_formatted_text(cell.paragraphs[0], cell_text)

            # Gray background for first column (field names) or header row
            if j == 0 or i == 0:
                set_cell_shading(cell, HEADER_GRAY)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = FONT_NAME

    doc.add_paragraph()


def is_uc_table(rows):
    """Detect if a table is a UC-style table (first cell empty, second cell has UC-XX)."""
    if len(rows) >= 2 and len(rows[0]) == 2:
        first_cell = rows[0][0].strip()
        second_cell = rows[0][1].strip()
        if first_cell == '' and ('UC-' in second_cell or 'C-0' in second_cell):
            return True
    # Also check if first column has field names like "Nombre", "Descripcion", etc
    if len(rows) >= 3 and len(rows[0]) == 2:
        field_names = {r[0].strip().lower() for r in rows if len(r) >= 2}
        uc_fields = {'nombre', 'descripcion', 'precondicion', 'postcondicion', 'secuencia normal'}
        if len(field_names & uc_fields) >= 2:
            return True
    return False


def convert_md_to_docx():
    """Main conversion function."""
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Configure styles
    configure_styles(doc)

    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Add cover page
    add_cover_page(doc)

    idx = 0
    skip_cover = True  # Skip the first lines that form the cover page in MD

    # Skip until after first "---" separator (cover page content)
    while idx < len(lines):
        line = lines[idx].strip()
        if line == '---' and skip_cover:
            idx += 1
            # Find second --- to skip Lista de Cambios header area
            break
        idx += 1

    # Now process the rest of the document
    # But first, let's skip the cover content we already added
    # Reset to process from "## Lista de Cambios"
    idx = 0
    in_cover = True

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip everything until "## Lista de Cambios"
        if in_cover:
            if stripped == '## Lista de Cambios':
                in_cover = False
                p = doc.add_heading('Lista de Cambios', level=2)
                idx += 1
                continue
            idx += 1
            continue

        # Skip empty lines
        if not stripped:
            idx += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            idx += 1
            continue

        # Headings
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                # Don't re-add "Lista de Cambios"
                h = doc.add_heading(level=min(level, 3))
                h.clear()
                add_formatted_text(h, text)
                for run in h.runs:
                    run.font.name = FONT_NAME
                    run.font.color.rgb = RGBColor(0, 0, 0)
                idx += 1
                continue

        # Tables
        if stripped.startswith('|'):
            rows, next_idx = parse_table_rows(lines, idx)
            if rows:
                if is_uc_table(rows):
                    add_uc_table(doc, rows)
                else:
                    add_table_to_doc(doc, rows)
            idx = next_idx
            continue

        # List items
        if stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            add_formatted_text(p, text)
            idx += 1
            continue

        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            add_formatted_text(p, text)
            idx += 1
            continue

        # Italic block (like *[Pendiente...]*)
        if stripped.startswith('*[') or stripped.startswith('*Pendiente'):
            p = doc.add_paragraph()
            text = stripped.strip('*')
            run = p.add_run(text)
            set_run_font(run, italic=True)
            idx += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        idx += 1

    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")


if __name__ == '__main__':
    convert_md_to_docx()
